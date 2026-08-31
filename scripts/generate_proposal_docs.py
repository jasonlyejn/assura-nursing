import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable

def create_docx(filename):
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    primary_color = RGBColor(13, 58, 84)    # #0D3A54
    accent_color = RGBColor(2, 132, 199)    # #0284C7
    dark_color = RGBColor(30, 41, 59)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("ASSURA NURSING CARE")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = primary_color

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Strategic Business Partner & Investment Pitch Proposal\nTransforming Home Healthcare Through Clinical Tech & 80/20 Economics")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = accent_color

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 1: Executive Summary
    h1 = doc.add_heading("1. Executive Summary", level=1)
    h1.runs[0].font.color.rgb = primary_color
    p = doc.add_paragraph(
        "Assura Nursing Care is a tech-enabled, hospital-grade home healthcare platform operating in Malaysia. "
        "We deploy licensed Registered Nurses (LJM certified) equipped with a proprietary clinical operating system "
        "featuring cloud MEWS scoring, digital MAR medication charting, wound healing staging, tube expiry tracking, "
        "and digital procedure consent signatures."
    )
    p.runs[0].font.color.rgb = dark_color

    # Section 2: Market Problem & Opportunity
    h1 = doc.add_heading("2. The Market Problem & Opportunity", level=1)
    h1.runs[0].font.color.rgb = primary_color
    
    doc.add_paragraph(
        "• Silver Economy Boom: Malaysia is officially an aging nation (7%+ above 65) reaching super-aged status by 2030, "
        "with over 3.5 million seniors requiring chronic disease and post-stroke home care.\n"
        "• Hospital Bed Shortages: Private and public tertiary hospitals face severe bed constraints and seek trusted home care partners for early discharge.\n"
        "• Fragmented Home Nursing Sector: Traditional agencies rely on pen-and-paper charts, freelance caregivers without clinical tracking, and zero transparency for doctors and families."
    )

    # Section 3: Core Competitive Moats
    h1 = doc.add_heading("3. Assura's 8 Unfair Competitive Moats", level=1)
    h1.runs[0].font.color.rgb = primary_color

    moats = [
        ("1. Hospital-Grade Clinical System", "Built-in MEWS vital score calculations, automated vital escalation alerts, and digital MAR drug administration."),
        ("2. Scalable 80/20 Unit Economics", "Nurses receive 80% take-home pay; Assura retains 20% platform margin with zero fixed nurse salary overhead."),
        ("3. 72-Hour Doctor Referral Link", "Temporary PIN-protected portal for hospital specialists to review vitals and wound healing without app installation."),
        ("4. Digital Procedure Consent Signing", "Touchscreen signature capture before invasive procedures, shielding the business from legal liability."),
        ("5. Bedside Consumables Auto-Billing", "Automated consumption billing on Clock-Out across ~80 hospital-grade catalog items at 40-60% margins."),
        ("6. 1-Tap Emergency SOS Beacon", "Instant 999 dialing and GPS dispatch coordinate generation for acute emergencies."),
        ("7. Family WhatsApp Arrival & Care Alerts", "Automated arrival alerts on Clock-In and comprehensive vitals summaries on Clock-Out."),
        ("8. Ultra-Light Native Mobile Apps", "4.1MB signed Android releases with biometric PIN security and offline cloud sync.")
    ]

    for title, desc in moats:
        p = doc.add_paragraph()
        r1 = p.add_run(f"• {title}: ")
        r1.font.bold = True
        r1.font.color.rgb = primary_color
        r2 = p.add_run(desc)
        r2.font.color.rgb = dark_color

    # Section 4: Financial Model Table
    h1 = doc.add_heading("4. Revenue Channels & Financial Economics", level=1)
    h1.runs[0].font.color.rgb = primary_color

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_titles = ['Service Stream', 'Client Pricing', 'Platform Share', 'Margin %']
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0D3A54"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)

    data = [
        ("Procedure Home Visits", "RM 120 - RM 350 / visit", "RM 24 - RM 70 / visit", "20% + Consumables"),
        ("Long-Term Nursing Shifts", "RM 250 - RM 650 / shift", "RM 50 - RM 130 / shift", "20% Commission"),
        ("Medical Consumables", "Catalog unit pricing", "Direct markup profit", "40% - 60% Gross Margin"),
        ("Hospital Discharge Partner", "Contract referral packages", "Volume recurring revenue", "High recurring LTV"),
    ]

    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    # Section 5: Partnership Models
    h1 = doc.add_heading("5. Strategic Partnership & Investment Models", level=1)
    h1.runs[0].font.color.rgb = primary_color

    doc.add_paragraph(
        "1. Equity / Strategic Growth Investor: Capital injection for statewide nurse recruitment and marketing.\n"
        "2. Clinical & Medical Partners (Specialists & Hospitals): Patient referral partnerships with revenue sharing.\n"
        "3. City / Regional Franchise Operator: Turnkey licensing of the Assura platform to operate in new cities."
    )

    doc.save(filename)
    print(f"✓ Saved Word Proposal to {filename}")


def create_pdf(filename):
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40,
    )
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=22,
        textColor=colors.HexColor('#0D3A54'),
        alignment=1,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=11,
        textColor=colors.HexColor('#0284C7'),
        alignment=1,
        spaceAfter=15,
    )
    h1_style = ParagraphStyle(
        'SectionH1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13,
        textColor=colors.HexColor('#0D3A54'),
        spaceBefore=12,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        'BodyDark',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        textColor=colors.HexColor('#1E293B'),
        leading=14,
        spaceAfter=6,
    )

    story = []
    story.append(Paragraph("ASSURA NURSING CARE", title_style))
    story.append(Paragraph("Strategic Business Partner &amp; Investment Pitch Deck", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#0D3A54'), spaceAfter=12))

    story.append(Paragraph("1. Executive Summary", h1_style))
    story.append(Paragraph(
        "<b>Assura Nursing Care</b> is a tech-enabled, hospital-grade home healthcare platform operating in Malaysia. "
        "We deploy licensed Registered Nurses (LJM certified) equipped with a proprietary clinical operating system "
        "featuring cloud MEWS scoring, digital MAR medication charting, wound healing staging, tube expiry tracking, "
        "and digital procedure consent signatures.",
        body_style
    ))

    story.append(Paragraph("2. Market Opportunity &amp; Silver Economy", h1_style))
    story.append(Paragraph(
        "• <b>Silver Economy Boom:</b> Malaysia is officially an aging nation (7%+ above 65) reaching super-aged status by 2030.<br/>"
        "• <b>Hospital Bed Crunch:</b> Tertiary hospitals face severe bed constraints and seek trusted home care partners for early discharge.<br/>"
        "• <b>Fragmented Home Nursing:</b> Traditional agencies rely on pen-and-paper charts and unverified caregivers.",
        body_style
    ))

    story.append(Paragraph("3. Core Competitive Moats", h1_style))
    story.append(Paragraph(
        "1. <b>Hospital-Grade Clinical System:</b> Cloud MEWS vital score calculations, automated vital escalation alerts.<br/>"
        "2. <b>Scalable 80/20 Unit Economics:</b> Nurses receive 80% take-home pay; Assura retains 20% platform margin.<br/>"
        "3. <b>72-Hour Doctor Referral Link:</b> Temporary PIN-protected portal for hospital specialists to review vitals.<br/>"
        "4. <b>Digital Procedure Consent:</b> Touchscreen signature capture before invasive procedures.<br/>"
        "5. <b>Bedside Consumables Auto-Billing:</b> Automated consumption billing on Clock-Out (~80 items, 40-60% margins).<br/>"
        "6. <b>1-Tap Emergency SOS Beacon:</b> Instant 999 dialing and GPS dispatch coordinate generation.",
        body_style
    ))

    story.append(Paragraph("4. Revenue Model &amp; Financial Economics", h1_style))
    table_data = [
        [Paragraph('<b>Service Stream</b>', body_style), Paragraph('<b>Client Pricing</b>', body_style), Paragraph('<b>Platform Share</b>', body_style), Paragraph('<b>Margin</b>', body_style)],
        ['Procedure Visits', 'RM 120 - RM 350 / visit', 'RM 24 - RM 70 / visit', '20% + Consumables'],
        ['Long-Term Shifts', 'RM 250 - RM 650 / shift', 'RM 50 - RM 130 / shift', '20% Commission'],
        ['Medical Consumables', 'Catalog unit pricing', 'Direct markup profit', '40% - 60% Margin'],
        ['Hospital Discharge', 'Contract referral packages', 'Volume recurring revenue', 'High recurring LTV'],
    ]
    t = Table(table_data, colWidths=[130, 140, 130, 110])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0D3A54')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(t)
    story.append(Spacer(1, 10))

    story.append(Paragraph("5. Partnership Models", h1_style))
    story.append(Paragraph(
        "• <b>Equity / Strategic Investor:</b> Capital injection for regional marketing and nurse acquisition.<br/>"
        "• <b>Clinical Partners (Doctors / Clinics):</b> Referral partnership with revenue sharing.<br/>"
        "• <b>Regional City Franchise:</b> Turnkey licensing to expand Assura to new states.",
        body_style
    ))

    doc.build(story)
    print(f"✓ Saved PDF Proposal to {filename}")

if __name__ == '__main__':
    create_docx('c:/assura/Assura_Business_Partner_Proposal.docx')
    create_pdf('c:/assura/Assura_Business_Partner_Proposal.pdf')
    create_docx('c:/assura/website/Assura_Business_Partner_Proposal.docx')
    create_pdf('c:/assura/website/Assura_Business_Partner_Proposal.pdf')
