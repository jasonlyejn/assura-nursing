import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

# Set Matplotlib font to support Chinese characters without square boxes (tofu)
plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'SimSun', 'DengXian', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, HRFlowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Register Windows TrueType Font for ReportLab CJK
font_regular = 'C:/Windows/Fonts/msyh.ttc' if os.path.exists('C:/Windows/Fonts/msyh.ttc') else 'C:/Windows/Fonts/simhei.ttf'
font_bold = 'C:/Windows/Fonts/msyhbd.ttc' if os.path.exists('C:/Windows/Fonts/msyhbd.ttc') else font_regular

try:
    pdfmetrics.registerFont(TTFont('CJK-Font', font_regular))
    pdfmetrics.registerFont(TTFont('CJK-Bold', font_bold))
    has_cjk_font = True
except Exception as e:
    print("Warning: Could not register CJK font for PDF:", e)
    has_cjk_font = False

CHART_DIR = 'c:/assura/scripts/charts'
os.makedirs(CHART_DIR, exist_ok=True)

def generate_charts():
    # --- CHART 1: Market Demand & Silver Economy ---
    fig, ax1 = plt.subplots(figsize=(7.2, 3.8), dpi=300)
    years = ['2022', '2024', '2026', '2028', '2030 (Super-Aged)']
    seniors_mil = [2.4, 2.9, 3.5, 4.3, 5.2]
    market_val_rm = [650, 920, 1450, 2300, 3600]

    color1 = '#0D3A54'
    color2 = '#10B981'

    ax1.set_title('Malaysia Home Healthcare & Silver Economy Growth (2022 - 2030)\n马来西亚居家医疗与银发经济规模增长趋势', fontsize=11, fontweight='bold', pad=12, color='#0D3A54')
    bars = ax1.bar(years, seniors_mil, color=color1, width=0.45, label='Senior Population 65+ (Millions / 百万人)')
    ax1.set_ylabel('Senior Population (Millions / 百万人)', color=color1, fontweight='bold', fontsize=9)
    ax1.tick_params(axis='y', labelcolor=color1)
    ax1.set_ylim(0, 6.5)

    for bar in bars:
        yval = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2.0, yval + 0.15, f'{yval}M', ha='center', va='bottom', fontsize=8.5, fontweight='bold', color=color1)

    ax2 = ax1.twinx()
    line = ax2.plot(years, market_val_rm, color=color2, marker='o', linewidth=2.5, markersize=7, label='Market Size (RM Mil / 百万令吉)')
    ax2.set_ylabel('Market Size (RM Millions / 百万令吉)', color=color2, fontweight='bold', fontsize=9)
    ax2.tick_params(axis='y', labelcolor=color2)
    ax2.set_ylim(0, 4200)

    for i, txt in enumerate(market_val_rm):
        ax2.annotate(f'RM {txt}M', (years[i], market_val_rm[i] + 120), ha='center', fontsize=8.5, fontweight='bold', color=color2)

    fig.tight_layout()
    chart1_path = os.path.join(CHART_DIR, 'chart1_market.png')
    fig.savefig(chart1_path)
    plt.close(fig)

    # --- CHART 2: Tiered Unit Economics & Multi-Stream Margin Waterfall ---
    fig, ax = plt.subplots(figsize=(7.2, 3.6), dpi=300)
    categories = [
        'Gross Client Order\n(单笔综合消费 100%)',
        'Field Staff Payout\n(护士/看护报酬 56-65%)',
        'Procedure Net Commission\n(专科护理佣金 36%)',
        'Consumables & Rental Net\n(耗材与辅具租赁净利 62%)',
        'Total Assura Net Margin\n(平台综合单笔净利润 43%)'
    ]
    amounts = [300, 170, 80, 50, 130]
    colors_list = ['#0D3A54', '#64748B', '#0284C7', '#F59E0B', '#10B981']

    bars = ax.barh(categories, amounts, color=colors_list, height=0.52)
    ax.set_title('Assura Tiered Unit Economics: Service Commission + Consumables Markup\n阶梯式单笔订单盈利模型 (35%平台抽佣 + 60%+耗材辅具净利润)', fontsize=10.5, fontweight='bold', pad=12, color='#0D3A54')
    ax.set_xlabel('Amount in MYR (RM / 令吉)', fontweight='bold', fontsize=9, color='#334155')
    ax.set_xlim(0, 350)

    for bar in bars:
        w = bar.get_width()
        ax.text(w + 5, bar.get_y() + bar.get_height()/2.0, f'RM {w:.0f}', ha='left', va='center', fontsize=9, fontweight='bold', color='#0F172A')

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    chart2_path = os.path.join(CHART_DIR, 'chart2_economics.png')
    fig.savefig(chart2_path)
    plt.close(fig)

    # --- CHART 3: Prudent 6-Year Regional Expansion & High-EBITDA Trajectory ---
    fig, ax = plt.subplots(figsize=(7.2, 3.6), dpi=300)
    years_proj = [
        'Years 1 – 2 (Penang & Kedah)\n第1–2年 (槟城北马深耕基地)',
        'Years 3 – 4 (Ipoh & Care Center)\n第3–4年 (怡保扩张+实体护理中心)',
        'Years 5 – 6 (Klang Valley / KL)\n第5–6年 (雪隆与吉隆坡大都会区)'
    ]
    revenue_proj = [1.65, 5.20, 12.80] # In RM Millions
    net_profit = [0.56, 1.95, 5.10]    # In RM Millions (34% - 40% net EBITDA margin)

    x = np.arange(len(years_proj))
    width = 0.35

    rects1 = ax.bar(x - width/2, revenue_proj, width, label='Gross Platform GMV (RM Mil / 总营业额)', color='#0D3A54')
    rects2 = ax.bar(x + width/2, net_profit, width, label='Net EBITDA Profit (RM Mil / 净利润 34-40%)', color='#10B981')

    ax.set_title('Assura 6-Year Prudent Expansion & Net EBITDA Projections (2025 - 2030)\n稳健六年阶梯式财务模型 · 高净利润率与稳健现金流', fontsize=10.5, fontweight='bold', pad=12, color='#0D3A54')
    ax.set_xticks(x)
    ax.set_xticklabels(years_proj, fontweight='bold', fontsize=8)
    ax.set_ylabel('RM (Millions / 百万令吉)', fontweight='bold', fontsize=9, color='#334155')
    ax.set_ylim(0, 15)
    ax.legend(frameon=True, loc='upper left')

    for rect in rects1:
        h = rect.get_height()
        ax.text(rect.get_x() + rect.get_width()/2., h + 0.25, f'RM {h:.2f}M', ha='center', va='bottom', fontsize=8, fontweight='bold', color='#0D3A54')

    for rect in rects2:
        h = rect.get_height()
        ax.text(rect.get_x() + rect.get_width()/2., h + 0.25, f'RM {h:.2f}M', ha='center', va='bottom', fontsize=8, fontweight='bold', color='#10B981')

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    chart3_path = os.path.join(CHART_DIR, 'chart3_projection.png')
    fig.savefig(chart3_path)
    plt.close(fig)

    print("✓ High-resolution CJK-compatible charts generated.")
    return chart1_path, chart2_path, chart3_path


def create_bilingual_docx(filename, c1, c2, c3):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    primary = RGBColor(13, 58, 84)     # #0D3A54
    accent = RGBColor(2, 132, 199)     # #0284C7
    dark = RGBColor(30, 41, 59)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ASSURA NURSING CARE · 安舒居家医疗护理")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = primary

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("STRATEGIC BUSINESS PARTNER & INVESTMENT PROPOSAL\n商业合作伙伴招募与高净利阶梯式投资商业计划书")
    r2.font.name = "Microsoft YaHei"
    r2.font.size = Pt(11.5)
    r2.font.bold = True
    r2.font.color.rgb = accent

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Tiered Take-Rate Model · Prudent 6-Year Expansion · High EBITDA Margins | 阶梯式高抽成·六年中长期稳健扩张·高净利模型")
    r3.font.name = "Microsoft YaHei"
    r3.font.size = Pt(9.5)
    r3.font.italic = True
    r3.font.color.rgb = dark

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 1: Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Investment Thesis / 发展愿景与投资亮点", level=1)
    h1.runs[0].font.color.rgb = primary
    doc.add_paragraph(
        "【English】 Assura Nursing Care is Malaysia's premier tech-enabled private home nursing and clinical caregiver platform. "
        "We deploy 100% licensed Registered Nurses (LJM certified) and trained professional caregivers (with flexible male/female selection) equipped with a proprietary cloud clinical operating system—including "
        "real-time MEWS early warning scoring, digital MAR medication charting, wound staging timelines, tube expiry tracking, and "
        "digital procedure consent signing.\n\n"
        "Rather than burning capital on risky expansion, Assura adheres to a prudent, cash-flow-positive, high-margin phased roadmap: "
        "establishing Penang and the Northern corridor in Years 1–2, expanding to Ipoh (Perak silver retirement corridor) and launching our physical Care Center in Years 3–4, and scaling to Klang Valley in Years 5–6. "
        "By optimizing our platform with a Tiered Take-Rate (35% on procedures, 25-30% on shifts) and capturing 50%+ gross margin on medical consumables and equipment rentals, Assura generates superior EBITDA margins (34%–40%) and rapid payback for investors.\n\n"
        "【中文】 Assura Nursing Care（安舒居家医疗护理）是马来西亚领先的科技赋能型医疗级居家护理与专业看护平台。我们全员采用具备大马护士局（LJM）执照的注册护士与受训专业医护人员（支持男/女护士自由指定），"
        "依托自主研发的云端临床管理系统——涵盖 MEWS 生命体征预警、电子用药图表（MAR）、伤口愈合追踪、导管到期预警及电子知情同意书。\n\n"
        "我们坚决拒绝盲目烧钱的急躁扩张，而是采取【稳扎稳打、注重高净利润与现金流】的阶梯式战略：第1–2年深耕槟城北马基地；第3–4年进军怡保（大马银发退休重镇）并落地实体护理康复中心；第5–6年成熟扩张至雪隆大都会区。"
        "通过优化【阶梯式抽成模型】（专科操作抽成 35-40%，长期排班抽成 25-30%）以及【耗材辅具零售 50%+ 纯利】，Assura 实现了 34%–40% 的高净利润率（EBITDA），为投资人提供极具确定性的投资回报周期。"
    )

    # Section 2: Live Demo Accounts for Investors & Partners
    h1 = doc.add_heading("2. Live Interactive Demo Accounts / 实机演示账号与体验入口", level=1)
    h1.runs[0].font.color.rgb = primary
    doc.add_paragraph(
        "Investors and partners can experience the live platform instantly using the dedicated demonstration credentials below:\n"
        "投资人与合作伙伴可直接使用以下演示账号登录体验完整系统："
    )

    demo_table = doc.add_table(rows=1, cols=4)
    demo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    demo_table.style = 'Table Grid'
    hdr = demo_table.rows[0].cells
    hdr_names = ['Role / 角色', 'Portal URL / 访问链接', 'Login Phone / 手机号', 'Security PIN / 密码']
    for i, name in enumerate(hdr_names):
        hdr[i].text = name
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0D3A54"/>')
        hdr[i]._tc.get_or_add_tcPr().append(shd)

    demos = [
        ("Admin / 院长主管", "https://staff.assuranursing.com", "0120001111", "1234"),
        ("Field Staff Nurse / 上门护士", "https://staff.assuranursing.com", "0120002222", "1234"),
        ("Patient & Family / 病患家属", "https://assuranursing.com/portal.html", "0120003333", "1234"),
        ("Attending Doctor / 主治医生", "https://assuranursing.com/doctor.html", "Case ID: demo_case_001", "PIN: 1234"),
    ]
    for d_row in demos:
        rc = demo_table.add_row().cells
        for i, val in enumerate(d_row):
            rc[i].text = val
            rc[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 3: Market Analysis & Chart 1
    h1 = doc.add_heading("3. The Market Opportunity & Silver Economy / 银发经济与市场机遇", level=1)
    h1.runs[0].font.color.rgb = primary
    doc.add_paragraph(
        "• Rapidly Aging Nation (人口老龄化加剧): Malaysia officially transitioned into an aging society with over 3.5 million seniors requiring long-term chronic disease, post-stroke, and post-surgical home nursing.\n"
        "• Hospital Bed Constraints (公私立医院病床紧缺): Tertiary hospitals face chronic bed shortages and actively partner with home nursing providers for structured post-discharge care.\n"
        "• High Demand in Northern Region & Perak (北马及霹雳老龄化高需求): Penang and Ipoh represent Malaysia's highest density of retirees and senior citizens seeking premium private healthcare at home."
    )
    if os.path.exists(c1):
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        doc.add_picture(c1, width=Inches(6.2))
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 4: Core Selling Points & Advantages Matrix
    h1 = doc.add_heading("4. Assura's Core Advantages vs Traditional Agencies / 系统核心优势对比", level=1)
    h1.runs[0].font.color.rgb = primary

    comp_table = doc.add_table(rows=1, cols=3)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_table.style = 'Table Grid'
    chdr = comp_table.rows[0].cells
    c_names = ['Clinical Feature / 临床功能', 'Traditional Agency / 传统中介看护', 'Assura Tech Platform / Assura 智能平台']
    for i, name in enumerate(c_names):
        chdr[i].text = name
        chdr[i].paragraphs[0].runs[0].font.bold = True
        chdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0D3A54"/>')
        chdr[i]._tc.get_or_add_tcPr().append(shd)

    comps = [
        ("Staff Qualification (人员资质)", "Unverified freelance maid/helper (无牌看护)", "100% LJM Licensed Registered Nurses (大马注册护士)"),
        ("Staff Gender Choice (男女护士/护理师自选)", "Fixed or random assign, cannot choose (无法指定男女)", "Free choice of Male or Female Nurses & Caregivers (自由指定男/女护士或看护师)"),
        ("Vital Signs (生命体征质控)", "Pen and paper or not measured (纸张或不测)", "Cloud MEWS Early Warning Scoring (智能体征预警)"),
        ("Medication Safety (用药安全)", "Handwritten notes, high error risk (手写错乱)", "Digital MAR + Editable Administer Times (电子精准核对)"),
        ("Doctor Collaboration (医生联动)", "Zero doctor communication (与主治医生脱节)", "72-Hour PIN Doctor Share Link (72小时安全共享)"),
        ("Self-Service Roster (护士自助接单排班)", "Manual phone/chat coordination (人工繁琐)", "Interactive Roster Slot Request + Admin 1-Tap Approval (智能申请与审批)"),
        ("Legal Protection (法律风控)", "No written informed consent (无知情同意)", "Digital Touchscreen Consent Signature (电子免责签名)"),
        ("Emergency Response (紧急事件)", "Panic, delayed ambulance (家属慌乱延误)", "1-Tap SOS Beacon with GPS Dispatch (一键急救坐标)"),
    ]
    for c_row in comps:
        rc = comp_table.add_row().cells
        for i, val in enumerate(c_row):
            rc[i].text = val
            rc[i].paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 5: Strategic 6-Year Expansion Roadmap
    h1 = doc.add_heading("5. Prudent 6-Year Phased Expansion Strategy / 稳健务实的六年阶段性扩张规划", level=1)
    h1.runs[0].font.color.rgb = primary
    doc.add_paragraph(
        "【English】\n"
        "• Phase 1 (Years 1 – 2 — Penang & Northern Base / 槟城北马基地深耕):\n"
        "  Consolidating operations across Penang Island, Bukit Mertajam, Seberang Perai, Kulim, and Sungai Petani. We partner directly with hospital discharge lounges, establish brand authority, onboard 80+ licensed nurses, and prove unit economics with strong net profit and cash flow.\n\n"
        "• Phase 2 (Years 3 – 4 — Perak / Ipoh Expansion & Assura Care Center Hub / 怡保扩张与实体护理康复中心落地 - KIV):\n"
        "  Expanding into Ipoh (Malaysia's premier retirement hub with high senior demographics). Concurrently establishing the physical Assura Care Center & Senior Day Living Hub to serve as the regional inpatient step-down facility, consumables depot, and nurse training academy.\n\n"
        "• Phase 3 (Years 5 – 6 — Klang Valley & Greater KL Scaling / 雪隆与大吉隆坡都会区成熟复制):\n"
        "  Replicating our proven digital home nursing and care center model in Petaling Jaya, Subang Jaya, Kuala Lumpur, and Shah Alam—scaling into Malaysia's largest private healthcare market with mature operations.\n\n"
        "【中文】\n"
        "• 第一阶段（第 1–2 年 —— 槟城深耕与北马走廊）：全面覆盖槟岛、大山脚、威省、居林与双溪大年。直接对接各大公私立医院出院转介通道，沉淀 80+ 名注册护士人才池，以轻资产模式跑通单店模型并实现健康的正向现金流与 34% 净利润率。\n"
        "• 第二阶段（第 3–4 年 —— 进军怡保与霹雳中马走廊 + 实体专业护理中心落地 - KIV）：扩张至怡保（马来西亚老龄化人口最集中、养老需求最强劲的银发重镇）。同步设立 Assura 实体专业护理康复中心，作为全区域移动护士调度站、耗材仓储库及住院中风康复基地。\n"
        "• 第三阶段（第 5–6 年 —— 雪隆与大吉隆坡成熟复制）：将成熟的“移动居家护理 + 实体中心 + 数字化系统”标准模式，全面推向雪隆、八打灵再也、梳邦及吉隆坡高净值家庭市场。"
    )

    # Section 6: High Margin Drivers for Investors
    h1 = doc.add_heading("6. Optimized Tiered Take-Rate & High Margin Drivers / 阶梯式高抽佣与四大盈利支柱", level=1)
    h1.runs[0].font.color.rgb = primary
    doc.add_paragraph(
        "1. 🩺 Tiered Platform Take-Rate (阶梯式服务抽成 30%–40%):\n"
        "   • High-Skill Clinical Procedures (1-2 Hours): Client pays RM 220–260; Nurse receives RM 140–160 (high hourly earning); Company retains RM 80–100 (35%–40% Net Take Rate).\n"
        "   • 12h/24h Bedside Shifts: Client pays RM 240/12h; Caregiver receives RM 170; Company retains RM 70 (29% Net Take Rate).\n\n"
        "2. 📦 Bedside Medical Consumables & Equipment Rental (高毛利耗材销售与医疗设备租赁 50%+ 净利):\n"
        "   Direct supply of advanced wound dressings (hydrocolloids, silver foam), Foley catheters, and tracheostomy tubes during nurse visits, combined with monthly rental of hospital beds, oxygen concentrators, and suction machines.\n\n"
        "3. 💳 Upfront Retainer Subscription Plans (长期照护会员预付费订阅制):\n"
        "   1-month and 3-month recurring care packages paid upfront (RM 6,200–6,800/mo), securing guaranteed cash flow float with zero bad debt liabilities.\n\n"
        "4. 💻 Proprietary Unified Software Asset (零边际软件技术成本):\n"
        "   Our in-house clinical operating system powers both mobile visits and care center ward management with zero recurring 3rd-party SaaS software licensing costs."
    )

    # Section 7: Unit Economics & Chart 2
    h1 = doc.add_heading("7. Unit Economics & Multi-Stream Margin Waterfall / 单笔订单多重盈利模型", level=1)
    h1.runs[0].font.color.rgb = primary
    if os.path.exists(c2):
        doc.add_picture(c2, width=Inches(6.2))
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 8: 6-Year Financial Projections & Chart 3
    h1 = doc.add_heading("8. 6-Year Prudent Financial Projections & EBITDA Forecast / 六年中长期财务预测与净利润", level=1)
    h1.runs[0].font.color.rgb = primary
    if os.path.exists(c3):
        doc.add_picture(c3, width=Inches(6.2))
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 9: Partnership Models
    h1 = doc.add_heading("9. Strategic Partnership & Co-Founder Models / 战略合伙人加盟与共建模式", level=1)
    h1.runs[0].font.color.rgb = primary
    doc.add_paragraph(
        "1. 💎 Strategic Equity Investor (战略股权合伙人): Capital injection for nurse recruitment scaling, brand marketing, and Care Center physical facility acquisition.\n"
        "2. 🏥 Care Center Facility & Real Estate Partner (护理中心场地/地产合伙人): Co-developing the flagship physical Assura Care Center hub in Penang / Perak.\n"
        "3. 🩺 Medical Specialist & Hospital Partner (专科医生与医院联盟): Direct clinical referral partnership with prioritized discharge pipelines and revenue sharing.\n"
        "4. 🏢 Regional City Franchisee (城市分站合伙人/加盟): Turnkey licensing to expand Assura into new territories with complete operational playbooks."
    )

    # Section 10: Contact & Inquiries
    h1 = doc.add_heading("10. Direct Inquiries & Partnership Contact / 合作洽询与联系方式", level=1)
    h1.runs[0].font.color.rgb = primary
    doc.add_paragraph(
        "For strategic partnership discussions, equity allocations, clinical referrals, or Care Center co-development inquiries, please contact our management team directly:\n"
        "如需进一步探讨战略合伙、股权投资、实体护理中心共建、医院/医生转介合作或城市分站加盟，欢迎直接与我们联络：\n\n"
        "• 📱 WhatsApp / Mobile (官方 WhatsApp 专线): +60 12-206 4868 (https://wa.me/60122064868)\n"
        "• ✉️ Email (商务合作邮箱): admin@assuranursing.com / contact@assuranursing.com\n"
        "• 🌐 Official Website (官方网站): https://assuranursing.com\n"
        "• 🏥 Operations Base (北马总部基地): Bukit Mertajam, Penang, Malaysia (大山脚·槟城)"
    )

    doc.save(filename)
    print(f"✓ Saved Enriched Bilingual Word Proposal to {filename}")


def create_bilingual_pdf(filename, c1, c2, c3):
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    styles = getSampleStyleSheet()

    font_name = 'CJK-Font' if has_cjk_font else 'Helvetica'
    font_bold_name = 'CJK-Bold' if has_cjk_font else 'Helvetica-Bold'

    title_style = ParagraphStyle(
        'DocTitle', parent=styles['Normal'],
        fontName=font_bold_name, fontSize=17,
        textColor=colors.HexColor('#0D3A54'), alignment=1, spaceAfter=4,
    )
    sub_style = ParagraphStyle(
        'DocSub', parent=styles['Normal'],
        fontName=font_bold_name, fontSize=9.5,
        textColor=colors.HexColor('#0284C7'), alignment=1, spaceAfter=8,
    )
    h1_style = ParagraphStyle(
        'SecH1', parent=styles['Normal'],
        fontName=font_bold_name, fontSize=10.5,
        textColor=colors.HexColor('#0D3A54'), spaceBefore=8, spaceAfter=3,
    )
    body_style = ParagraphStyle(
        'BodyTxt', parent=styles['Normal'],
        fontName=font_name, fontSize=8,
        textColor=colors.HexColor('#1E293B'), leading=11.5, spaceAfter=3,
    )

    story = []
    story.append(Paragraph("ASSURA NURSING CARE · 安舒居家医疗护理", title_style))
    story.append(Paragraph("STRATEGIC BUSINESS PARTNER &amp; INVESTMENT PROPOSAL (商业计划书)", sub_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#0D3A54'), spaceAfter=6))

    story.append(Paragraph("1. Executive Summary &amp; Investment Thesis / 发展愿景与投资亮点", h1_style))
    story.append(Paragraph(
        "<b>Assura Nursing Care</b> is Malaysia's premier tech-enabled, hospital-grade home healthcare platform. "
        "We deploy licensed Registered Nurses (LJM certified) and trained professional caregivers equipped with a proprietary cloud clinical operating system "
        "(MEWS scoring, digital MAR medication charting, wound staging timelines, tube expiry tracking, and digital procedure consent). "
        "Assura executes a <b>prudent, cash-flow-positive phased roadmap</b>: consolidating Penang &amp; the Northern corridor in Years 1–2, expanding to Ipoh (Perak silver retirement corridor) and launching our physical Care Center in Years 3–4, and scaling to Klang Valley in Years 5–6. "
        "By implementing a <b>Tiered Take-Rate (35-40% on procedures, 25-30% on shifts)</b> and capturing 50%+ gross profit on consumables and rentals, Assura generates superior <b>34%–40% Net EBITDA margins</b>.<br/>"
        "<b>中文摘要：</b>Assura Nursing Care（安舒居家医疗护理）是马来西亚领先的科技赋能型医疗级居家护理平台。我们全员采用具备大马护士局（LJM）执照的注册护士与专业医护人员，依托自主研发的云端临床管理系统，采取【稳扎稳打、注重高净利润与现金流】的节奏：第1–2年深耕槟城北马，第3–4年进军怡保并落地实体护理中心，第5–6年成熟扩张至雪隆大都会区。通过优化阶梯式抽成与耗材辅具零售，实现 34%–40% 的高净利润率（EBITDA）。",
        body_style
    ))

    story.append(Paragraph("2. Live Interactive Demo Accounts / 实机演示账号与体验入口", h1_style))
    story.append(Paragraph("Investors &amp; partners can test the live system directly with these credentials:<br/>"
                           "• <b>Admin Portal:</b> https://staff.assuranursing.com (Phone: 0120001111 · PIN: 1234)<br/>"
                           "• <b>Staff Nurse:</b> https://staff.assuranursing.com (Phone: 0120002222 · PIN: 1234)<br/>"
                           "• <b>Patient Member:</b> https://assuranursing.com/portal.html (Phone: 0120003333 · PIN: 1234)<br/>"
                           "• <b>Doctor 72-Hr Link:</b> https://assuranursing.com/doctor.html (Case ID: demo_case_001 · PIN: 1234)", body_style))

    story.append(Paragraph("3. Market Opportunity &amp; Silver Economy / 银发经济与市场机遇", h1_style))
    if os.path.exists(c1):
        story.append(RLImage(c1, width=500, height=240))
        story.append(Spacer(1, 4))

    story.append(Paragraph("4. Core Clinical Advantages vs Traditional Agencies / 系统核心优势对比", h1_style))
    comp_data = [
        [Paragraph('<b>Feature / 功能</b>', body_style), Paragraph('<b>Traditional Agency / 传统中介</b>', body_style), Paragraph('<b>Assura Tech Platform / Assura 平台</b>', body_style)],
        ['Staff Qualification', 'Unverified maid/helper (无牌看护)', '100% LJM Licensed Nurses (注册护士)'],
        ['Staff Gender Choice', 'Random, cannot select (无法指定男女)', 'Male / Female Nurses & Caregivers (自由指定男女)'],
        ['Vital Signs Oversight', 'Pen/paper or no records (无体征图)', 'Cloud MEWS Early Warning Score (体征预警)'],
        ['Medication Safety', 'Handwritten notes, high error risk', 'Digital MAR + Editable Administer Time (电子核对)'],
        ['Self-Service Roster', 'Manual phone/chat coordination', 'Interactive Slot Request + Admin 1-Tap Approval (排班审批)'],
        ['Doctor Integration', 'Zero doctor communication', '72-Hour PIN Doctor Share Link (医生通道)'],
        ['Emergency SOS', 'Delayed ambulance, panic', '1-Tap SOS Beacon with GPS Dispatch (一键急救)'],
    ]
    ct = Table(comp_data, colWidths=[120, 170, 210])
    ct.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0D3A54')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(ct)
    story.append(Spacer(1, 4))

    story.append(Paragraph("5. Prudent 6-Year Phased Expansion / 稳健务实的六年阶段性扩张规划", h1_style))
    story.append(Paragraph(
        "• <b>Phase 1 (Years 1 – 2 / 槟城北马深耕):</b> Consolidating Penang Island, Bukit Mertajam, Seberang Perai, Kulim &amp; Sg Petani. 80+ nurses, strong positive cash flow (34% Net EBITDA margin).<br/>"
        "• <b>Phase 2 (Years 3 – 4 / 怡保扩张与实体护理中心落地 - KIV):</b> Expanding to Ipoh (Perak silver retirement corridor) + opening the flagship physical Assura Care Center &amp; Senior Day Living Hub.<br/>"
        "• <b>Phase 3 (Years 5 – 6 / 雪隆大都会区成熟复制):</b> Scaling the proven playbook into Klang Valley, PJ, Subang &amp; KL with mature operations and brand authority.",
        body_style
    ))

    story.append(Paragraph("6. Tiered Take-Rate &amp; High-Margin Revenue Pillars / 阶梯式高抽佣与四大盈利支柱", h1_style))
    story.append(Paragraph(
        "1. <b>Tiered Platform Take-Rate (阶梯抽成 30–40%):</b> 35-40% on 1-2h procedures (RM 80–100/visit) + 25-30% on 12h shifts.<br/>"
        "2. <b>Consumables &amp; Equipment Rental (耗材与设备租赁 50%+ 净利):</b> Bedside wound dressings, catheters + monthly rentals of hospital beds &amp; oxygen concentrators.<br/>"
        "3. <b>Upfront Retainer Subscriptions (预付费会员包月制):</b> Recurring monthly packages paid in advance (zero bad debt, strong cash flow float).<br/>"
        "4. <b>Unified Cloud Architecture (一套系统零边际成本):</b> Zero 3rd-party software licensing overhead for both mobile care and inpatient center.",
        body_style
    ))

    story.append(Paragraph("7. Unit Economics &amp; Margin Waterfall / 单笔订单多重盈利模型", h1_style))
    if os.path.exists(c2):
        story.append(RLImage(c2, width=500, height=220))
        story.append(Spacer(1, 4))

    story.append(Paragraph("8. 6-Year Financial Forecast &amp; Net EBITDA / 六年财务预测与净利润 (34-40% Margin)", h1_style))
    if os.path.exists(c3):
        story.append(RLImage(c3, width=500, height=225))
        story.append(Spacer(1, 4))

    story.append(Paragraph("9. Partnership &amp; Investment Models / 战略合伙人加盟模式", h1_style))
    story.append(Paragraph(
        "• <b>Equity / Strategic Investor (股权投资):</b> Capital injection for nurse recruitment scaling, brand marketing, and Care Center physical facility acquisition.<br/>"
        "• <b>Facility &amp; Real Estate Partner (护理中心场地合伙人):</b> Co-developing the physical Assura Care Center hub in Penang / Perak.<br/>"
        "• <b>Clinical Partners (医生/医院联盟):</b> Patient referral partnerships with prioritized pipelines and revenue sharing.<br/>"
        "• <b>Regional City Franchise (城市分站加盟):</b> Turnkey licensing to expand Assura to new states with proven operational playbooks.",
        body_style
    ))

    story.append(Paragraph("10. Direct Inquiries &amp; WhatsApp Contact / 合作洽询与联系方式", h1_style))
    story.append(Paragraph(
        "• <b>WhatsApp 专线:</b> +60 12-206 4868 (https://wa.me/60122064868)<br/>"
        "• <b>Email 邮箱:</b> admin@assuranursing.com<br/>"
        "• <b>Website 官网:</b> https://assuranursing.com · <b>Operations Base:</b> Bukit Mertajam, Penang",
        body_style
    ))

    doc.build(story)
    print(f"✓ Saved Enriched Bilingual PDF Proposal to {filename}")


if __name__ == '__main__':
    c1, c2, c3 = generate_charts()
    create_bilingual_docx('c:/assura/Assura_Business_Partner_Proposal_Bilingual.docx', c1, c2, c3)
    create_bilingual_pdf('c:/assura/Assura_Business_Partner_Proposal_Bilingual.pdf', c1, c2, c3)
    create_bilingual_docx('c:/assura/website/Assura_Business_Partner_Proposal_Bilingual.docx', c1, c2, c3)
    create_bilingual_pdf('c:/assura/website/Assura_Business_Partner_Proposal_Bilingual.pdf', c1, c2, c3)
